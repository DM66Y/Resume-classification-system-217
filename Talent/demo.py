# coding:utf-8

import numpy as np
import jieba
import gensim
import docx
from keras.models import load_model
from keras import backend as K

# 加载维基中文词向量模型
path = './简历分类测试/wiki.zh.text.model'


# 自定义metrics（这个是训练模型时的准确率函数，在这里没有用到但是必须加上，否则报错）
def precision(y_true, y_pred):
    # Calculates the precision
    true_positives = K.sum(K.round(K.clip(y_true * y_pred, 0, 1)))
    predicted_positives = K.sum(K.round(K.clip(y_pred, 0, 1)))
    precision = true_positives / (predicted_positives + K.epsilon())
    return precision


def recall(y_true, y_pred):
    # Calculates the recall
    true_positives = K.sum(K.round(K.clip(y_true * y_pred, 0, 1)))
    possible_positives = K.sum(K.round(K.clip(y_true, 0, 1)))
    recall = true_positives / (possible_positives + K.epsilon())
    return recall

# 中文结巴分词(同时导入词表中未出现的专有名词)
# 下面的这个information变量就是我们得到的待分类的信息（通过用户输入获取），下面只是我的一个样例
jieba.load_userdict('./简历分类测试/dmy_userwords.txt')

# 载入已经训好的h5模型
work_dir = r'./简历分类测试/jianli.h5'
print('开始加载模型')
model = load_model(work_dir, custom_objects={'precision': precision, 'recall': recall})
print('load success')



def classify(file_name):
    file = docx.Document(file_name)
    print("段落数:" + str(len(file.paragraphs)))  # 输出段落数
    file_word = docx.Document()
    information = ''
    for para in file.paragraphs:
        information = information + para.text
    # print(information)
    # information = docx.Document("E:/DMY important/段铭杨毕业设计/A-GCNN/测试文档.docx")
    # information = "项目经验人力资源000000管理系统项目时间：2019年01月-2019年02月项目简介： 使用技术及语言：SSM框架/Java 开发工具：Eclipse/STS 开发人员： 独立完成 项目描述：管理员对部门、员工、职位、培训、招聘、奖惩、考勤的CRUD；游客浏览招聘，对自己简历的CRUD，投递；员工查看自己的信息、培训、公司通讯录、奖惩、薪资和打卡。本项目结构上分为表现层、业务层和数据访问层，层次间的依赖关系自下到上。业务层封装业务流程，为适应业务的变更，每一业务模块均有专门的接口及实现类。 项目业绩： 项目收获：只是离现在最近的小项目，花费半个多月的时间完成。功能繁琐，但我没有放弃，虚心请教一些大佬，在此我也由衷感谢他们。对于代码的编写，一定要字斟句酌，一个不起眼的问题，就会导致系统BUG。程序员也是在代码的海洋里磨炼起来的！"
    a_cut = jieba.cut(information)
    a = '/'.join(a_cut)
    after_cut = a.split('/')
    # print(after_cut)

    for i in range(3):
        for i in after_cut:
            if len(i) > 2:
                after_cut.remove(i)
        # print(after_cut)

    final_cut = []
    for item in after_cut:
        if item != ' ' and item != '  ' and item != '':
            final_cut.append(item)
    # print(final_cut)

    # 去噪，导入中文停用词表
    file_stop = r'./简历分类测试/dmy_stopwords.txt'
    stop = []
    standard_stop = []
    final_information = []
    with open(file_stop, 'r', encoding='utf-8-sig') as f:
        lines = f.readlines()
        for line in lines:
            lline = line.strip()
            stop.append(lline)

    for i in range(0, len(stop)):
        for word in stop[i].split():
            standard_stop.append(word)

    for i in final_cut:
        if i not in standard_stop:
            final_information.append(i)
    # print(final_information)

    # 将得到的文本向量化
    model_1 = gensim.models.Word2Vec.load(path)
    X = []
    sentence = []
    word_vec = []
    zero_array = [0] * 256
    word_vec = np.zeros([256, ], dtype=np.float64)
    for j in range(len(final_information)):
        word_vec = list(model_1[final_information[j]])
        sentence.append(word_vec)
    for j in range(128 - len(final_information)):
        sentence.append(zero_array)
    X.append(sentence)
    X = np.array(X)

    predict = model.predict(X)
    print(predict)

    # 上面的predict就是最终的分类结果，是一个数组，其中分值最大对应的类别标签就为判定的类别
    if predict[0][0] > predict[0][1] and predict[0][0] > predict[0][2] and predict[0][0] > predict[0][3] and predict[0][
        0] > predict[0][4]:
        return 'JAVA工程师'
    if predict[0][1] > predict[0][2] and predict[0][1] > predict[0][3] and predict[0][1] > predict[0][4] and predict[0][
        1] > predict[0][0]:
        return '技术总监'
    if predict[0][2] > predict[0][3] and predict[0][2] > predict[0][4] and predict[0][2] > predict[0][1] and predict[0][
        2] > predict[0][0]:
        return 'Web工程师'
    if [0][3] > predict[0][4] and predict[0][3] > predict[0][2] and predict[0][3] > predict[0][1] and predict[0][3] > \
            predict[0][0]:
        return '大数据工程师'
    if predict[0][4] > predict[0][3] and predict[0][4] > predict[0][2] and predict[0][4] > predict[0][0] and predict[0][
        4] > predict[0][1]:
        return '算法工程师'


if __name__ == "__main__":
    print(classify('test_dir/1.docx'))
